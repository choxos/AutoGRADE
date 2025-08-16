"""
Summary of Findings table generator following Core GRADE format
Based on Core GRADE 6 guidance for presenting evidence in summary tables
"""
import json
import time
from typing import Dict, List, Optional, Tuple
from django.conf import settings
from django.template.loader import render_to_string
import anthropic

from ..models import (
    GRADEProject, Outcome, SummaryOfFindingsTable, 
    GRADEAssessment, PlainLanguageStatement, AIAnalysisSession
)


class SoFTableGenerator:
    """
    Generates Summary of Findings tables following GRADE methodology
    """
    
    def __init__(self):
        if not settings.ANTHROPIC_API_KEY:
            raise ValueError("ANTHROPIC_API_KEY not configured in settings")
        self.client = anthropic.Anthropic(api_key=settings.ANTHROPIC_API_KEY)
        self.model = "claude-3-5-sonnet-20241022"
    
    def generate_sof_table(self, project: GRADEProject, user) -> SummaryOfFindingsTable:
        """
        Generate a complete Summary of Findings table for the project
        """
        # Get or create SoF table
        sof_table, created = SummaryOfFindingsTable.objects.get_or_create(
            project=project,
            defaults={
                'title': f"Summary of findings: {project.title}",
                'population': project.population,
                'intervention': project.intervention,
                'comparison': project.comparison,
                'generated_by': user
            }
        )
        
        if not created:
            # Update existing table
            sof_table.population = project.population
            sof_table.intervention = project.intervention
            sof_table.comparison = project.comparison
            sof_table.generated_by = user
            sof_table.save()
        
        return sof_table
    
    def generate_enhanced_sof_with_ai(self, project: GRADEProject, user) -> Dict:
        """
        Generate enhanced SoF table content with AI assistance
        """
        start_time = time.time()
        
        # Gather data for SoF table
        sof_data = self._gather_sof_data(project)
        
        # Create AI prompt for enhanced SoF generation
        prompt = self._create_sof_prompt(project, sof_data)
        
        try:
            # Call Claude Sonnet
            response = self.client.messages.create(
                model=self.model,
                max_tokens=4000,
                temperature=0.1,
                messages=[{
                    "role": "user",
                    "content": prompt
                }]
            )
            
            ai_response = response.content[0].text
            processing_time = time.time() - start_time
            
            # Parse AI response
            enhanced_sof_data = self._parse_sof_response(ai_response)
            
            # Create or update SoF table
            sof_table = self.generate_sof_table(project, user)
            sof_table.ai_generation_data = {
                'ai_response': ai_response,
                'enhanced_data': enhanced_sof_data
            }
            sof_table.save()
            
            # Log AI session
            self._log_ai_session(
                project, 
                'sof_generation', 
                sof_data, 
                ai_response, 
                processing_time, 
                user
            )
            
            return {
                'sof_table': sof_table,
                'enhanced_data': enhanced_sof_data,
                'success': True
            }
            
        except Exception as e:
            # Log error session
            self._log_ai_session(
                project, 
                'sof_generation', 
                sof_data, 
                None, 
                time.time() - start_time, 
                user, 
                error=str(e)
            )
            raise Exception(f"Enhanced SoF generation failed: {str(e)}")
    
    def _gather_sof_data(self, project: GRADEProject) -> Dict:
        """
        Gather all data needed for SoF table generation
        """
        # Get critical and important outcomes
        outcomes = project.outcomes.filter(importance__gte=4).order_by('-importance', 'name')
        
        sof_data = {
            'project': {
                'title': project.title,
                'population': project.population,
                'intervention': project.intervention,
                'comparison': project.comparison,
            },
            'outcomes': []
        }
        
        for outcome in outcomes:
            outcome_data = {
                'name': outcome.name,
                'description': outcome.description,
                'importance': outcome.importance,
                'importance_category': 'Critical' if outcome.importance >= 7 else 'Important',
                'outcome_type': outcome.outcome_type,
                'measurement_scale': outcome.measurement_scale,
                'time_frame': outcome.time_frame,
                'minimal_important_difference': outcome.minimal_important_difference,
                
                # Effect estimates
                'relative_effect': outcome.relative_effect,
                'relative_effect_type': outcome.relative_effect_type,
                'confidence_interval_lower': outcome.confidence_interval_lower,
                'confidence_interval_upper': outcome.confidence_interval_upper,
                'baseline_risk': outcome.baseline_risk,
                'intervention_risk': outcome.intervention_risk,
                'risk_difference': outcome.risk_difference,
                'risk_difference_ci_lower': outcome.risk_difference_ci_lower,
                'risk_difference_ci_upper': outcome.risk_difference_ci_upper,
                
                # Study data
                'number_of_studies': outcome.number_of_studies,
                'total_participants': outcome.total_participants,
                
                # GRADE assessment
                'grade_assessment': None,
                'plain_language': None
            }
            
            # Add GRADE assessment if available
            try:
                grade_assessment = outcome.grade_assessment
                outcome_data['grade_assessment'] = {
                    'final_certainty': grade_assessment.final_certainty,
                    'starting_certainty': grade_assessment.starting_certainty,
                    'risk_of_bias': grade_assessment.risk_of_bias,
                    'inconsistency': grade_assessment.inconsistency,
                    'indirectness': grade_assessment.indirectness,
                    'imprecision': grade_assessment.imprecision,
                    'publication_bias': grade_assessment.publication_bias,
                    'rating_down_reasons': self._get_rating_down_reasons(grade_assessment)
                }
            except GRADEAssessment.DoesNotExist:
                pass
            
            # Add plain language statement if available
            try:
                plain_lang = outcome.plain_language
                outcome_data['plain_language'] = {
                    'statement': plain_lang.statement,
                    'certainty_description': plain_lang.certainty_description
                }
            except PlainLanguageStatement.DoesNotExist:
                pass
            
            sof_data['outcomes'].append(outcome_data)
        
        return sof_data
    
    def _get_rating_down_reasons(self, assessment: GRADEAssessment) -> List[str]:
        """
        Get list of reasons for rating down certainty
        """
        reasons = []
        
        if assessment.risk_of_bias_rating_down > 0:
            level = "serious" if assessment.risk_of_bias_rating_down == 1 else "very serious"
            reasons.append(f"Risk of bias ({level})")
        
        if assessment.inconsistency_rating_down > 0:
            level = "serious" if assessment.inconsistency_rating_down == 1 else "very serious"
            reasons.append(f"Inconsistency ({level})")
        
        if assessment.indirectness_rating_down > 0:
            level = "serious" if assessment.indirectness_rating_down == 1 else "very serious"
            reasons.append(f"Indirectness ({level})")
        
        if assessment.imprecision_rating_down > 0:
            level = "serious" if assessment.imprecision_rating_down == 1 else "very serious"
            reasons.append(f"Imprecision ({level})")
        
        if assessment.publication_bias_rating_down > 0:
            level = "serious" if assessment.publication_bias_rating_down == 1 else "very serious"
            reasons.append(f"Publication bias ({level})")
        
        return reasons
    
    def _create_sof_prompt(self, project: GRADEProject, sof_data: Dict) -> str:
        """
        Create the prompt for Claude Sonnet to enhance SoF table content
        """
        prompt = f"""
        You are an expert in GRADE methodology and Summary of Findings table creation. Please help enhance and validate the following Summary of Findings table data according to Core GRADE principles.

        CORE GRADE SUMMARY OF FINDINGS TABLE REQUIREMENTS:
        
        Based on Core GRADE 6, a SoF table should include:
        1. Clear identification of population, intervention, comparison, and setting
        2. Key outcomes ordered by importance (critical first, then important)
        3. For each outcome:
           - Number of studies and participants
           - Relative effect estimates with confidence intervals
           - Absolute effect estimates for different risk groups when relevant
           - Certainty of evidence rating (high/moderate/low/very low)
           - Reasons for rating down (if applicable)
           - Plain language summary of findings
        
        CURRENT SOF DATA:
        {json.dumps(sof_data, indent=2)}
        
        ENHANCEMENT TASKS:
        
        Please provide enhancements and validation in the following JSON format:
        
        {{
            "table_title": "Improved title following GRADE format",
            "bibliography_citation": "Suggested citation format for this systematic review",
            
            "population_enhancement": {{
                "description": "Enhanced population description",
                "setting": "Healthcare setting details",
                "characteristics": "Key population characteristics"
            }},
            
            "intervention_comparison_enhancement": {{
                "intervention_details": "Enhanced intervention description",
                "comparison_details": "Enhanced comparison description",
                "intervention_format": "Formatted intervention name for table header",
                "comparison_format": "Formatted comparison name for table header"
            }},
            
            "outcomes_enhancements": [
                {{
                    "outcome_name": "Original outcome name",
                    "enhanced_description": "Clearer outcome description",
                    "measurement_details": "Enhanced measurement description",
                    "time_frame_clarity": "Clarified time frame",
                    "absolute_effect_interpretation": "Interpretation of absolute effects",
                    "clinical_significance": "Clinical significance assessment",
                    "plain_language_improvement": "Improved plain language statement",
                    "certainty_rationale": "Clear explanation of certainty rating"
                }}
            ],
            
            "table_formatting": {{
                "column_headers": ["List of appropriate column headers"],
                "footnotes": ["Important footnotes for the table"],
                "abbreviations": {{"abbreviation": "definition"}},
                "quality_indicators": "Symbols or indicators for certainty levels"
            }},
            
            "overall_assessment": {{
                "strengths": ["Key strengths of the evidence"],
                "limitations": ["Key limitations to highlight"],
                "applicability": "Comments on applicability of findings",
                "certainty_summary": "Overall certainty of evidence summary"
            }},
            
            "recommendations": {{
                "data_gaps": ["Important data gaps identified"],
                "additional_analyses": ["Suggested additional analyses"],
                "presentation_improvements": ["Ways to improve table presentation"]
            }}
        }}
        
        Focus on:
        1. Ensuring compliance with Core GRADE 6 standards
        2. Improving clarity and clinical interpretability
        3. Identifying any missing or unclear information
        4. Enhancing plain language statements
        5. Ensuring appropriate certainty ratings and rationales
        """
        
        return prompt
    
    def _parse_sof_response(self, ai_response: str) -> Dict:
        """
        Parse the AI response into structured SoF enhancement data
        """
        try:
            # Try to extract JSON from the response
            json_start = ai_response.find('{')
            json_end = ai_response.rfind('}') + 1
            
            if json_start == -1 or json_end == 0:
                raise ValueError("No JSON found in response")
            
            json_str = ai_response[json_start:json_end]
            sof_enhancement_data = json.loads(json_str)
            
            return sof_enhancement_data
            
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON in AI response: {str(e)}")
        except Exception as e:
            raise ValueError(f"Error parsing AI response: {str(e)}")
    
    def _log_ai_session(self, project: GRADEProject, session_type: str, input_data: Dict, 
                       ai_response: str, processing_time: float, user, error: str = None):
        """
        Log AI analysis session for debugging and improvement
        """
        AIAnalysisSession.objects.create(
            project=project,
            session_type=session_type,
            input_data=input_data,
            ai_response={'response': ai_response} if ai_response else {},
            processing_time=processing_time,
            error_occurred=bool(error),
            error_message=error or '',
            created_by=user
        )
    
    def generate_html_table(self, project: GRADEProject) -> str:
        """
        Generate HTML representation of the SoF table
        """
        sof_data = self._gather_sof_data(project)
        
        # Template context
        context = {
            'project': project,
            'outcomes': sof_data['outcomes'],
            'critical_outcomes': [o for o in sof_data['outcomes'] if o['importance'] >= 7],
            'important_outcomes': [o for o in sof_data['outcomes'] if 4 <= o['importance'] < 7],
        }
        
        # This would use a Django template - for now return a basic structure
        html = self._generate_basic_html_table(context)
        
        return html
    
    def _generate_basic_html_table(self, context: Dict) -> str:
        """
        Generate basic HTML table structure
        """
        project = context['project']
        outcomes = context['outcomes']
        
        html = f"""
        <div class="sof-table-container">
            <h2>Summary of Findings</h2>
            <div class="sof-header">
                <p><strong>Population:</strong> {project.population}</p>
                <p><strong>Intervention:</strong> {project.intervention}</p>
                <p><strong>Comparison:</strong> {project.comparison}</p>
            </div>
            
            <table class="sof-table table table-bordered">
                <thead class="table-header">
                    <tr>
                        <th rowspan="2">Outcome</th>
                        <th rowspan="2">Study Results and Measurements</th>
                        <th colspan="2">Absolute Effect Estimates</th>
                        <th rowspan="2">Certainty of Evidence</th>
                        <th rowspan="2">Plain Language Summary</th>
                    </tr>
                    <tr>
                        <th>{project.comparison}</th>
                        <th>{project.intervention}</th>
                    </tr>
                </thead>
                <tbody>
        """
        
        for outcome in outcomes:
            certainty_class = f"certainty-{outcome['grade_assessment']['final_certainty'].replace('_', '-')}" if outcome.get('grade_assessment') else 'certainty-unknown'
            
            html += f"""
                    <tr class="outcome-row">
                        <td class="outcome-name">
                            <strong>{outcome['name']}</strong>
                            <br><small>{outcome['time_frame']}</small>
                        </td>
                        <td class="study-results">
                            <div class="effect-estimate">
                                {outcome['relative_effect_type'] or 'Effect'}: {outcome['relative_effect'] or 'Not reported'}
                                {f"(95% CI: {outcome['confidence_interval_lower']} to {outcome['confidence_interval_upper']})" if outcome['confidence_interval_lower'] else ''}
                            </div>
                            <div class="study-info">
                                Based on {outcome['number_of_studies'] or 'N/A'} studies with {outcome['total_participants'] or 'N/A'} participants
                            </div>
                        </td>
                        <td class="control-group">
                            {f"{outcome['baseline_risk']} per 1000" if outcome['baseline_risk'] else 'Not reported'}
                        </td>
                        <td class="intervention-group">
                            {f"{outcome['intervention_risk']} per 1000" if outcome['intervention_risk'] else 'Not reported'}
                            {f"<br>({outcome['risk_difference']} fewer/more per 1000)" if outcome['risk_difference'] else ''}
                        </td>
                        <td class="certainty {certainty_class}">
                            <div class="certainty-rating">
                                {outcome['grade_assessment']['final_certainty'].replace('_', ' ').title() if outcome.get('grade_assessment') else 'Not assessed'}
                            </div>
                            {f"<div class='rating-reasons'>{', '.join(outcome['grade_assessment']['rating_down_reasons'])}</div>" if outcome.get('grade_assessment', {}).get('rating_down_reasons') else ''}
                        </td>
                        <td class="plain-language">
                            {outcome['plain_language']['statement'] if outcome.get('plain_language') else 'No plain language summary available'}
                        </td>
                    </tr>
            """
        
        html += """
                </tbody>
            </table>
        </div>
        
        <style>
            .sof-table-container {
                font-family: Arial, sans-serif;
                margin: 20px 0;
            }
            .sof-table {
                width: 100%;
                border-collapse: collapse;
                margin: 20px 0;
            }
            .sof-table th, .sof-table td {
                border: 1px solid #ddd;
                padding: 12px;
                text-align: left;
                vertical-align: top;
            }
            .table-header {
                background-color: #f8f9fa;
                font-weight: bold;
            }
            .outcome-name {
                background-color: #f1f3f4;
                width: 15%;
            }
            .certainty-high { background-color: #d4edda; }
            .certainty-moderate { background-color: #fff3cd; }
            .certainty-low { background-color: #f8d7da; }
            .certainty-very-low { background-color: #f5c6cb; }
            .certainty-rating {
                font-weight: bold;
                text-transform: capitalize;
            }
            .rating-reasons {
                font-size: 0.9em;
                font-style: italic;
                margin-top: 4px;
            }
            .effect-estimate {
                font-weight: bold;
                margin-bottom: 4px;
            }
            .study-info {
                font-size: 0.9em;
                color: #666;
            }
        </style>
        """
        
        return html
    
    def export_to_docx(self, project: GRADEProject, file_path: str = None) -> str:
        """
        Export SoF table to Microsoft Word format
        """
        try:
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            
            # Add title
            title = doc.add_heading(f'Summary of Findings: {project.title}', 0)
            
            # Add PICO information
            doc.add_heading('Study Question', level=1)
            p = doc.add_paragraph()
            p.add_run('Population: ').bold = True
            p.add_run(project.population)
            p.add_run('\nIntervention: ').bold = True
            p.add_run(project.intervention)
            p.add_run('\nComparison: ').bold = True
            p.add_run(project.comparison)
            
            # Add outcomes table
            sof_data = self._gather_sof_data(project)
            
            if sof_data['outcomes']:
                doc.add_heading('Summary of Findings Table', level=1)
                
                # Create table
                table = doc.add_table(rows=1, cols=6)
                table.style = 'Table Grid'
                
                # Header row
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Outcome'
                hdr_cells[1].text = 'Studies/Participants'
                hdr_cells[2].text = 'Effect Estimate'
                hdr_cells[3].text = 'Absolute Effects'
                hdr_cells[4].text = 'Certainty'
                hdr_cells[5].text = 'Plain Language Summary'
                
                # Add outcome rows
                for outcome in sof_data['outcomes']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = outcome['name']
                    row_cells[1].text = f"{outcome['number_of_studies'] or 'N/A'} studies, {outcome['total_participants'] or 'N/A'} participants"
                    
                    effect_text = f"{outcome['relative_effect_type'] or ''}: {outcome['relative_effect'] or 'Not reported'}"
                    if outcome['confidence_interval_lower']:
                        effect_text += f" (95% CI: {outcome['confidence_interval_lower']} to {outcome['confidence_interval_upper']})"
                    row_cells[2].text = effect_text
                    
                    absolute_text = ""
                    if outcome['risk_difference']:
                        absolute_text = f"{outcome['risk_difference']} per 1000 difference"
                    row_cells[3].text = absolute_text or 'Not reported'
                    
                    certainty_text = outcome['grade_assessment']['final_certainty'].replace('_', ' ').title() if outcome.get('grade_assessment') else 'Not assessed'
                    row_cells[4].text = certainty_text
                    
                    row_cells[5].text = outcome['plain_language']['statement'] if outcome.get('plain_language') else 'No summary available'
            
            # Save document
            if not file_path:
                file_path = f"sof_table_{project.id}.docx"
            
            doc.save(file_path)
            return file_path
            
        except ImportError:
            raise Exception("python-docx is required for Word export functionality")
        except Exception as e:
            raise Exception(f"Error exporting to Word: {str(e)}")
    
    def validate_sof_completeness(self, project: GRADEProject) -> Dict:
        """
        Validate completeness of SoF table data
        """
        validation_results = {
            'complete': True,
            'warnings': [],
            'errors': [],
            'missing_data': []
        }
        
        # Check basic project information
        if not project.population:
            validation_results['errors'].append("Population not defined")
        if not project.intervention:
            validation_results['errors'].append("Intervention not defined")
        if not project.comparison:
            validation_results['errors'].append("Comparison not defined")
        
        # Check outcomes
        outcomes = project.outcomes.filter(importance__gte=4)
        if not outcomes.exists():
            validation_results['errors'].append("No important or critical outcomes defined")
        
        for outcome in outcomes:
            outcome_issues = []
            
            # Check effect estimates
            if not outcome.relative_effect:
                outcome_issues.append("Missing relative effect estimate")
            if not outcome.number_of_studies:
                outcome_issues.append("Missing number of studies")
            if not outcome.total_participants:
                outcome_issues.append("Missing participant count")
            
            # Check GRADE assessment
            try:
                assessment = outcome.grade_assessment
            except GRADEAssessment.DoesNotExist:
                outcome_issues.append("Missing GRADE assessment")
            
            # Check plain language statement
            try:
                plain_lang = outcome.plain_language
            except PlainLanguageStatement.DoesNotExist:
                outcome_issues.append("Missing plain language statement")
            
            if outcome_issues:
                validation_results['missing_data'].append({
                    'outcome': outcome.name,
                    'issues': outcome_issues
                })
        
        # Set overall completeness
        if validation_results['errors'] or validation_results['missing_data']:
            validation_results['complete'] = False
        
        return validation_results
